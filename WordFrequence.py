import docx
import glob
from pattern import en
import re
from nltk.corpus import wordnet as wn
import nltk

###nltk.sent_tokenize(text)  段落->句子
# import nltk
#
# sent_text = nltk.sent_tokenize(text) # this gives us a list of sentences
# # now loop over each sentence and tokenize it separately
# for sentence in sent_text:
#     tokenized_text = nltk.word_tokenize(sentence)
#     tagged = nltk.pos_tag(tokenized_text)
#     print(tagged)


###获取单词的基本形式，包括副词转形容词
def get_basic_form(wd):
    word=wd[0]
    tag=wd[1]
    print('wd=', wd)
    wnl = nltk.stem.WordNetLemmatizer()
    result=word
    if tag.startswith('VB'):
        result=wnl.lemmatize(word, 'v')
    elif tag.startswith("NN"):
        result=wnl.lemmatize(word, 'n')
    elif tag.startswith('JJ'):
        result=wnl.lemmatize(word, 'a')
    elif tag.startswith("R"):
        result=adverb2adject(word)[0]
    return result


###解析句子，由于pattern.en.parse不好用，错误较多，故使用nltk的，并增加副词转形容词
def parse_sentence(sentence):
    result=[]
    tokens=nltk.word_tokenize(sentence)

    for word, tag in nltk.pos_tag(tokens):  ###
        bas=get_basic_form((word,tag))
        result.append((word.lower(), tag, bas.lower()))

    return result

###副词转形容词
def adverb2adject(word):
    possible_adj = []
    for ss in nltk.corpus.wordnet.synsets(word):
        for lemmas in ss.lemmas():  # all possible lemmas

            for ps in lemmas.pertainyms():  # all possible pertainyms
                possible_adj.append(ps.name())
    if 'more' in word or 'most' in word :
        return [word]
    if not possible_adj:
        return [word]
    return possible_adj
##------------------------------------
def is_chinese_sentence(sentence):
    num_ch=0
    for ch in sentence:
        if is_Chinese(ch):
            num_ch+=1
    if num_ch>len(sentence)*0.50:
        return True
    else:
        return  False

###判断单个字符位汉字
def is_Chinese(ch):
    if '\u4e00' <= ch <= '\u9fff':
        return True
    else:
        return False

#--删除小括号和里面的内容（中英文都如此，一般括号里的中文是解释，英文是简写）
def remove_words_in_brackets(sentence):
    tt=re.findall(r'([^\(\)]*)\([^)]*\)([^\(\)]*)', sentence)
    if not tt:
        return sentence

    result=''
    for ll in tt:
        for mm in ll:
            result=result+' '+ mm
    return result

##--处理一个句子，count_word_frequence函数太长，就把它移除来
def process_sentence(sentence, words):
    ''' words 的格式
    {
    book:{'freqs'(（总次数）):9, 'NNS':['books', 5(某个词性的次数),[sentences](所有例句)]    },
    }
    '''
    ''' words 的格式
    {
    book:{'freqs'(（总次数）):9, 'NNS':['books', {}(单词的翻译、音标和读音), 5(某个词性的次数),[sentences](所有例句)]    },
    }
    '''
    ll = en.parse(sentence, relations=True, lemmata=True)
    # print('ll=', ll)
    for word in ll.split()[0]:
        basic_form = word[-1]   ####wd是单词的基础型
        this_word=word[0].lower()       #原单词
        word_tag=word[1]    ##词性
        if not basic_form.isalpha():
            continue
        if basic_form in words:  ###该单词出现过
            words[basic_form]['freqs']+=1  ##总频率加1

            if word[1] in words[basic_form]:  ##该单词的该词性出现过
                words[basic_form][word[1]][1] += 1  ###频率加1
                words[basic_form][word[1]][2].append(sentence)
            else: ##该单词的该词性没出现过
                words[basic_form][word[1]] = [word[0].lower(), 1, [sentence]]
        else: ###该单词从未出现过
            words[basic_form] = {'freqs':1,  word_tag:[this_word, 1, [sentence] ]}

    return True
# ------计算词频-----
def count_word_frequence(docs):
    words={}
    setences_all=[]

    for doc in docs:
        for paragraph in doc.paragraphs:
            if len(paragraph.text)< 10:  ##太短，段落里面没需要的内容
                continue
            if is_chinese_sentence(paragraph.text):
                # print('中文语句:', paragraph.text)
                continue
            sentences=re.findall(r'([a-zA-Z][^\.\?!]*[\.\?!])', paragraph.text)
            setences_all.extend(sentences)

    for sentence in setences_all:
        if len(sentence)< 10:  ####太短，肯定不是句子
            continue
        if is_chinese_sentence(sentence):
            continue

        sentence=remove_words_in_brackets(sentence)
        process_sentence2(sentence, words)

    return words

###
##--第2种方法：  处理一个句子，  count_word_frequence函数太长，就把它移除来
def process_sentence2(sentence, words):
    ''' words 的格式
    {
    book:{'freqs'(（总次数）):9, 'NNS':['books', 5(某个词性的次数),[sentences](所有例句)]    },
    }
    '''
    ''' words 的格式
    {
    book:{'freqs'(（总次数）):9, 'NNS':['books', {}(单词的翻译、音标和读音), 5(某个词性的次数),[sentences](所有例句)]    },
    }
    '''
    ll =parse_sentence(sentence)
    # print('ll=', ll)
    for word in ll:
        basic_form = word[-1]   ####wd是单词的基础型
        this_word=word[0].lower()       #原单词
        word_tag=word[1]    ##词性
        if not basic_form.isalpha():
            continue
        if basic_form in words:  ###该单词出现过
            words[basic_form]['freqs']+=1  ##总频率加1

            if word[1] in words[basic_form]:  ##该单词的该词性出现过
                words[basic_form][word[1]][1] += 1  ###频率加1
                words[basic_form][word[1]][2].append(sentence)
            else: ##该单词的该词性没出现过
                words[basic_form][word[1]] = [word[0].lower(), 1, [sentence]]
        else: ###该单词从未出现过
            words[basic_form] = {'freqs':1,  word_tag:[this_word, 1, [sentence] ]}

    return True


def count_by_cixin(words_all):    #分别计算动词、名词、形容词的高频词汇，并排序
    cixin_list=['CC', 'CD', 'DT', 'EX', 'FW', 'IN', 'JJ', 'JJR', 'JJS', 'LS', 'MD', 'NN', 'NNS', 'NNP', 'NNPS', 'PDT', 'POS', 'PRP', 'PRP$', 'RB', 'RBR', 'RBS', 'RP', 'SYM', 'TO', 'UH', 'VB', 'VBD', 'VBG', 'VBN', 'VBP', 'VBZ', 'WDT', 'WP', 'WP$', 'WRB']
    cixin_data={}
    for item in cixin_list:
        cixin_data[item]=[0,set()]
    others=0

    for word in words_all:
        for cixin in words_all[word]:
            if cixin=='freqs':
                continue
            # print('word=',words_all[word])
            if cixin in cixin_list:
                cixin_data[cixin][0]+=1
                cixin_data[cixin][1].add(words_all[word][cixin][0])
            else:
                print(cixin)
                others+=1
    return cixin_data

def read_sentences_from_docs(docs):
    setences_all=[]
    for doc in docs:
        for paragraph in doc.paragraphs:
            if len(paragraph.text) < 10:  ##太短，段落里面没需要的内容
                continue
            if is_chinese_sentence(paragraph.text):
                # print('中文语句:', paragraph.text)
                continue
            sentences = re.findall(r'([a-zA-Z][^\.\?!]*[\.\?!])', paragraph.text)
            setences_all.extend(sentences)
    return setences_all


if __name__=="__main__":
    dir0 = '高考真题'
    path_list = glob.glob(dir0 + '/*.docx')
    docs = []
    for p in path_list:
        if '$' in p:  ###略过临时文件
            continue
        doc = docx.Document(p)
        docs.append(doc)


    words_all = count_word_frequence(docs)
    import json
    with open('all_words.json', 'w', encoding='utf-8') as f:
        json.dump(words_all, f)


###--为所有单词查字典，并保存

###暂时注销
    # from crawl_dicts import get_translation_from_iciba
    # for word in words_all:
    #     for cixin in words_all[word]:
    #         if cixin=='freqs':
    #             continue
    #         tr=get_translation_from_iciba(words_all[word][cixin][0])
    #         words_all[word][cixin].insert(1, tr)
    #     print(words_all[word])
    # #write to file
    # with open('dicts_from_iciba.json', 'w', encoding='utf-8') as f:
    #     import json
    #     json.dump(words_all, f)