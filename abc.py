import re
import pystardict
from word_forms.word_forms import get_word_forms
import docx
from docx.shared import Pt
import datetime

maxNumber= 60

middle_wordList = []
senior_wordList = []
note_list = []



def is_check_word(word):
    for item in word_list:
        if(word == item):
            return (word , word_list[item])
    if word not in middle_wordList:
        return get_word_info(word)

def isAlpha(str1):
    if (str1>='a' and str1<='z' ) or  (str1>='a' and str1<='z'):
        return  True
    else:
        return False

def pharseLangdao(word):
    tt=word.split('\n')
    yinbiao=''
    fanyi=[]
    for item in tt:
        kk=item.lstrip(' ')
        # print('kk=',kk, isAlpha(kk[0]))
        if isAlpha(kk[0]):
            fanyi.append(kk)
        elif kk[0]=='*':
            yinbiao=kk[1:]
        else:
            break
    return (yinbiao, ','.join(fanyi))

def pharseQuick_eng(word):
    pass

def pharseDictItem( source ,item):  ##处理单个单词

    if source=='朗道':
        return pharseLangdao(item)
    elif source=='英中简明':
        return pharseQuick_eng(item)

def init_dicts():
    dict1 =( '朗道' ,'stardict-langdao-ec-gb-2.4.2/langdao-ec-gb')
    dict2 =('英中简明','stardict-quick_eng-zh_CN-2.4.2/quick_eng-zh_CN' )
    dict_path = [dict1, dict2]
    dicts=[]
    for i in range(0,len(dict_path)):
        x=[ dict_path[i][0], pystardict.Dictionary(dict_path[i][1])]
        dicts.append(x)
    return dicts

def get_yinbiao_fanyi( dicts, word):  ##获取音标和翻译

    for dic in dicts:
        dictItem = dic[1].get(word , 'no')
        # print('dictItem=', dictItem)
        if(dictItem!='no'):
            return pharseDictItem(dic[0] ,dictItem)

def gen_word_package(word):
    ll=get_word_forms(word)

    return ll

def removeTag(word):
    tags=[',', '.', '!', '?', ':', '-']
    for tag in tags:
        if tag in word:
            word=word.replace(tag, '')
    return word

def processLine(line):
    lineList = []
    index = 0
    for word in line.split():
        index += 1
        w2 = word
        w2=removeTag(w2)
        w2=w2.lower()
        checkRes = is_check_word(w2)
        lineList.append({word:checkRes})
    return lineList

def get_word_info(word_list, word):
    if word in word_list:
        return (word, word_list[word])
    else:
        wb=find_in_wordPackage(word)
        if wb:
            return (wb, word_list[wb])
        else:
            return None

def find_in_wordPackage(word):
    for item in word_package:
        for att in word_package[item]:
            if word in word_package[item][att] :
                return item
    return None

def screen_word(word):
    typeList = ['v' , 'n' , 'adj' , 'adv' , 'phr.' , 'pron' , 'prep' ]
    for type in typeList:
        if (word == type):
            return 'false'
    if(re.match(r'^[^a-zA-Z]' , word)):
        return 'false'
    return 'true'

def notes_words(endList):
    tempNotes = []
    tempNotes2 = []
    for oneDict in endList:
        for attr in oneDict:
            one = oneDict[attr]
            if oneDict[attr] is not None:
                if(len(tempNotes)>0):
                    if one not in tempNotes:
                        tempNotes.append(one)
                else:
                    tempNotes.append(one)

    for item in tempNotes:
        if item[1] is not None:
            tempNotes2.append(item)

    return tempNotes2

def complete_blank(word , say):
    pL = len(word)
    sL = len(say)
    lL = ((pL+1)*3-sL)/2
    lL = round(lL)
    print(word , '147')
    # allL = sL+lL*2
    # print(allL , '148')

    if lL <0:
        return say

    i = 0
    blankWord = ''
    while i < lL:
        blankWord = blankWord+'.'
        i+=1

    say = blankWord + say + blankWord
    return say

def get_a_line(wordList):
    showWord = ''
    lastWord = wordList[len(wordList)-1]
    for word in wordList:
        strlen = len(showWord)
        showWord = showWord + word +' '

        if(len(showWord)>maxNumber):
            showWord = showWord[:strlen]
            print(showWord, '165')
            set_other_line(showWord)
            showWord = word+' '
        elif(word == lastWord):
            set_other_line(showWord)
            showWord =  ' '

def set_other_line(showWord):
    print(showWord , 180)
    exit()
    showWordList = showWord.split(' ')
    sayWordList = []
    for word in showWordList:
        for item in endList:
            if word in item:
                wordMes = item[word]

                if (wordMes is not None and wordMes[1] is not None):
                    say = wordMes[1][0]
                else:
                    say = ''
                sayWord = complete_blank(word, say)
                sayWordList.append(sayWord)
                break


    write_a_line(showWordList , sayWordList)

def write_a_line(showArr , sayArr):
    showTxt = ''
    sayTxt = ''
    for show in showArr:
        showTxt = showTxt+show+' '
    for say in sayArr:
        sayTxt = sayTxt+say+' '
    write_to_word(sayTxt , showTxt)

def write_to_word(txt1 , txt2):
    paragraph1 = doc.add_paragraph()
    run1 = paragraph1.add_run(txt1)
    run1.font.size = Pt(8)
    paragraph2 = doc.add_paragraph()
    run2 = paragraph2.add_run(txt2)
    run2.font.size = Pt(16)
###
def get_sentences(path):
    with open(path, 'r', encoding='UTF-8') as f :
        tt=[]
        for ll in f.readlines():
            tt.append(ll)
        contents=''.join(tt)
        xx=contents.replace('\r', '').replace('\n', '')
        sentences=xx.split('.')
        return  sentences
# def intersection(obj1 , obj2):
def get_word_set(path):
    word_set=set()
    with open(path, 'r', encoding='UTF-8') as f:
        for line in f.readlines():
            for word in line.split():
                word = word.lower()
                word_set.add(word)
    return word_set
##------------init middle school vocabulary-------------------
def get_middle_wordset(path):
    word_set=set()
    with open(path,'r', encoding='utf-8') as f:
        for line in f.readlines():
            tt=re.findall(r'^\d{1,5}．([a-zA-Z]{1,10}).*$', line)
            if tt:
                word_set.add(tt[0].lower())

    return word_set
#----------init high school vocabulary------------------
def get_high_wordset(path):   ###不处理短语
    word_set=set()
    with open(path,'r', encoding='utf-8') as f:
        lines=f.readlines()
    r = r'^([a-zA-Z]{1,20})[\s]{1,5}\[.*\]'
    for line in lines:
        line=line.replace('[', ' [')

        tt= re.findall(r, line)
        if tt:
            if tt[0] in ['at', 'be', 'an', 'I', 'We', 'Set', 'man' ]:
                print('line=', line, 'word=', tt[0])
            word_set.add(tt[0].lower())

    return word_set
####-------------------main-------------------
middlePath = '初中词汇.txt'
highPath = '高中词汇(含短语).txt'
middle_wordset=get_middle_wordset(middlePath)
high_wordset=get_high_wordset(highPath)


word_list=dict()
word_package=dict()


cross=[]
for word in middle_wordset:
    if word in high_wordset:
        cross.append(word)
