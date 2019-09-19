###生成高考英语资料

import re
import pystardict
import docx
import datetime
# from pattern import en
import nltk
import json
from docx.shared import RGBColor
import difflib
import glob


def isAlpha(str1):
    if (str1 >= 'a' and str1 <= 'z') or (str1 >= 'a' and str1 <= 'z'):
        return True
    else:
        return False


'''  本地查找字典，意义不大，因为网上爬取单词的量比较小，解释也更好，还有音标啥的
def pharseLangdao(word):
    tt = word.split('\n')
    yinbiao = ''
    fanyi = []
    for item in tt:
        kk = item.lstrip(' ')
        # print('kk=',kk, isAlpha(kk[0]))
        if isAlpha(kk[0]):
            fanyi.append(kk)
        elif kk[0] == '*':
            yinbiao = kk[1:]
        else:
            break
    return (yinbiao, ','.join(fanyi))


def pharseQuick_eng(word):
    pass


def pharseDictItem(source, item):  ##处理单个单词

    if source == '朗道':
        return pharseLangdao(item)
    elif source == '英中简明':
        return pharseQuick_eng(item)


def init_dicts():
    dict1 = ('朗道', 'stardict-langdao-ec-gb-2.4.2/langdao-ec-gb')
    dict2 = ('英中简明', 'stardict-quick_eng-zh_CN-2.4.2/quick_eng-zh_CN')
    dict_path = [dict1, dict2]
    dicts = []
    for i in range(0, len(dict_path)):
        print(dict_path)
        x = [dict_path[i][0], pystardict.Dictionary(dict_path[i][1])]
        dicts.append(x)
    return dicts


def get_local_translation(dicts, word):  ##获取音标和翻译

    for dic in dicts:
        dictItem = dic[1].get(word, 'no')
        # print('dictItem=', dictItem)
        if (dictItem != 'no'):
            return pharseDictItem(dic[0], dictItem)
'''


###副词转形容词
def adverb2adject(word):
    possible_adj = []
    for ss in nltk.corpus.wordnet.synsets(word):
        for lemmas in ss.lemmas():  # all possible lemmas

            for ps in lemmas.pertainyms():  # all possible pertainyms
                possible_adj.append(ps.name())
    if 'more' in word or 'most' in word:
        return [word]
    if not possible_adj:
        return [word]

    nltk.edit_distance(word, possible_adj) / len(word)
    ratio = difflib.SequenceMatcher(None, word, possible_adj).ratio()
    if ratio < 0.5:  ###形容词和副词基本没关系，舍去
        possible_adj = word

    return possible_adj


###获取单词的基本形式，包括副词转形容词
def get_basic_form(wd):
    word = wd[0]
    tag = wd[1]
    # print('wd=', wd)
    wnl = nltk.stem.WordNetLemmatizer()
    result = word
    if tag.startswith('VB'):
        result = wnl.lemmatize(word, 'v')
    elif tag.startswith("NN"):
        result = wnl.lemmatize(word, 'n')
        if tag == "NNS" and result == word.lower():
            pass  ####未来再解决这些问题
        if tag == "NNPS" and result == word.lower():
            pass  ####未来再解决这些问题
    elif tag.startswith('JJ'):
        result = wnl.lemmatize(word, 'a')
    elif tag.startswith("R"):
        result = adverb2adject(word)[0]
    return result


###解析句子，由于pattern.en.parse不好用，错误较多，故使用nltk的，并增加副词转形容词
def parse_sentence(sentence):
    result = []
    tokens = nltk.word_tokenize(sentence)

    for word, tag in nltk.pos_tag(tokens):  ###
        bas = get_basic_form((word.lower(), tag))
        result.append((word, tag, bas.lower()))

    return result


# ----------------------------------------
def find_word_index(vocabulary, word):
    index = -1
    for dd in vocabulary:
        if word in dd:
            index = dd[0]
    return index


def get_word_entry(word):
    wd = word[1]
    fanyi = word[3]['translation']

    yinbiao = word[3]['yinbiao']
    yb = ''
    if '英' in yinbiao:
        yb = yinbiao['英']
    elif 'all' in yinbiao:
        yb = yinbiao['all']
    else:
        yb = ''
    fy = []
    for xx in fanyi:
        fy.append(xx + ' '.join(fanyi[xx]))
    return (wd, yb, fy)


###是否为单词(仅仅是由英文单词和-组成)
def isEnglishWord(word):
    if re.match(r'^[a-zA-Z][a-zA-Z\-]*[a-zA-Z]$', word):
        return True
    else:
        if word == 'I' or word == 'a':
            return True
        else:
            return False


# ----查找一个单词是否在单词列表里？---
def find_entry_index(vocabulary, word):
    index = -1
    for dd in vocabulary:
        if word in dd:
            index = dd[0]
    return index


###generate vocabulary for one reading, 包括高中词汇表和初中词汇表中的3音节及以上的词汇
def generate_vocabulary(paragraphs, high_words, middle_words, dicts):
    vocabulary = {}
    vocabulary['high'] = []
    vocabulary['extra'] = []
    vocabulary['blank'] = []

    high_index = 0
    extra_index = 0
    for paragraph in paragraphs:
        for sentence in paragraph:
            ll = parse_sentence(sentence)
            for word, word_tag, basic_word in ll:
                if not isEnglishWord(basic_word):
                    continue
                ###这些词性的单词一般都非常基础，忽略掉
                if word_tag in ['CC', 'CD', 'DT', 'IN', 'MD', 'WDT', 'WP', 'WP$', 'WRB']:
                    continue

                try:
                    translation = dicts[basic_word]
                    translation = dicts[word]
                except:
                    print('单词：' + basic_word + '  未收录！！')
                    vocabulary['blank'].append(word)
                    continue
                if basic_word in middle_words or word in middle_words:
                    if syllable_count(word) >= 3:  ###初中英语中的困难词汇
                        index = find_word_index(vocabulary['high'], word)
                        if index == -1:
                            high_index = high_index + 1
                            vocabulary['high'].append((high_index, word, dicts[word]))
                    continue
                elif basic_word in high_words or word in high_words:

                    index = find_word_index(vocabulary['high'], word)
                    if index == -1:
                        high_index = high_index + 1
                        vocabulary['high'].append((high_index, word, dicts[word]))

                else:
                    print('超纲词:', word, basic_word)
                    index = find_word_index(vocabulary['extra'], word)
                    if index == -1:
                        extra_index = extra_index + 1
                        vocabulary['extra'].append((extra_index, word, dicts[word]))

    return vocabulary


def get_entry(entries, word):
    for item in entries:
        if word in item:
            return item


def get_yinbiao2(entry):
    '''
    entry格式： （index, 'word',
    {
    'yinbiao':
        {
            '英': '[wɜ:d]',
            '美': '[wɜ:d]'
        },
    'translation':
        {

        }
    }）
    '''
    word = entry[0]
    yinbiao = entry[2]['yinbiao']
    if '英' in yinbiao:
        yb = yinbiao['英']
    elif 'all' in yinbiao:
        yb = yinbiao['all']
    else:
        yb = ''

    return yb


###get translation
def format_translation(entry):
    '''
    entry格式： （index, 'word',
    {
    'yinbiao':
        {
            '英': '[wɜ:d]',
            '美': '[wɜ:d]'
        },
    'translation':
        {

        }
    }）
    '''
    word = entry[0]
    ts = []
    translation = entry[2]['translation']
    for key in translation:
        xx = key + ' '.join(translation[key])
        ts.append(xx)
    return ts


###获取单词所属类型（）
def get_word_category(vocabulary, word):
    for item in vocabulary['high']:
        if word in item:
            return 'high'
    for item in vocabulary['extra']:
        if word in item:
            return 'extra'
    for item in vocabulary['blank']:
        if word in item:
            return 'blank'


###新的处理文章的方式
'''
step 1: 读取一篇文章，段落和句子分出来
step 2: 调用generate_vocabulary函数生成需要注音的单词列表，区分考纲单词和超纲单词
step 3:  生成带音标的文章

'''


def get_new_article(doc_src, high_words, middle_words, dicts):

    paragraphs = []
    word_cout = 0
    for paragraph in doc_src.paragraphs:
        paragraphs.append(nltk.sent_tokenize(paragraph.text))

    vocabulary = generate_vocabulary(paragraphs, high_words, middle_words, dicts)

    ##output to docx
    doc = docx.Document()

    ##输出词汇表（包含大纲词汇和超纲词汇），并可以修改样式
    p = doc.add_paragraph()
    p.add_run('大纲词汇').bold = True
    p = doc.add_paragraph()
    for entry in vocabulary['high']:
        p.add_run(str(entry[0]) + '. ')  ###这里可以设置样式
        p.add_run(entry[1])  ###这里可以设置样式
        run = p.add_run(get_yinbiao2(entry) + ' ')  ###这里可以设置样式
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        r = p.add_run(format_translation(entry))  ###这里可以设置样式
        r.add_break()

    p = doc.add_paragraph()
    p.add_run('超纲词汇').bold = True
    p = doc.add_paragraph()
    for entry in vocabulary['extra']:
        p.add_run('x' + str(entry[0]) + '. ')  ###这里可以设置样式
        p.add_run(entry[1])  ###这里可以设置样式
        p.add_run(get_yinbiao2(entry) + ' ')  ###这里可以设置样式
        run = p.add_run(format_translation(entry))  ###这里可以设置样式
        run.add_break()

    ###输出注好音过的文章内容-----

    for paragraph in paragraphs:
        p = doc.add_paragraph()
        p.add_run('    ')
        for sentence in paragraph:
            ll = parse_sentence(sentence)
            for word in ll:
                basic_form = word[-1]  ####wd是单词的基础型
                this_word = word[0]  # 原单词
                word_tag = word[1]  ##词性
                # print('word=', word)
                if not isEnglishWord(basic_form):
                    p.add_run(this_word + ' ')  ##不相关，直接输出
                    continue
                else:
                    word_cout += 1

                tt = get_word_category(vocabulary, this_word)
                if tt == 'high':
                    dd = p.add_run(this_word)
                    entry = get_entry(vocabulary['high'], this_word)

                    run = p.add_run(str(entry[0]))
                    run.font.superscript = True  ##单词的序号
                    run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

                    # 音标
                    yinbiao = p.add_run(get_yinbiao2(entry))
                    # font = yinbiao.font
                    yinbiao.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

                    p.add_run(' ')

                elif tt == 'extra':
                    print('extra：', this_word, basic_form)
                    dd = p.add_run(this_word)
                    entry = get_entry(vocabulary['extra'], this_word)
                    run = p.add_run('x' + str(entry[0]))  ##单词的序号
                    run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                    run.font.superscript = True

                    yinbiao = p.add_run(get_yinbiao2(entry))
                    yinbiao.font.color.rgb = RGBColor(0x3f, 0x2c, 0x36)

                    p.add_run(' ')
                elif tt == 'blank':
                    dd = p.add_run(this_word + ' ')
                else:
                    dd = p.add_run(this_word + ' ')

    return (doc, word_cout, vocabulary)


##estimate difficulty of a reading through its vocabulary
def evaluate_difficulty(doc_src, vocabulary, method):
    '''
    methods=1(Gunning Fog Index), 2(The Flesch Formula), 3(Power Sumner Kearl Formula)
The Flesch Formula: （数值越大越容易）
0-29 Very Difficult Post Graduate
30-49 Difficult College
50-59 Fairly Difficult High School
60-69 Standard 8th to 9th grade
70-79 Fairly Easy 7th grade
80-89 Easy 5th to 6th grade
90-100 Very Easy 4th to 5th grade
    '''
    paragraphs=[]
    for paragraph in doc_src.paragraphs:
        paragraphs.append(nltk.sent_tokenize(paragraph.text))
    number_of_sentence=0
    number_of_syllables=0
    number_of_words=0
    number_of_big_words=0
    for paragraph in paragraphs:

        number_of_sentence+=len(paragraph)

        for sentence in paragraph:
            ll = parse_sentence(sentence)
            for word in ll:
                if not isEnglishWord(word[-1]):
                    continue
                number_of_words+=1
                syc=syllable_count(word[-1])
                number_of_syllables+=syc
                if syc>=3:
                    number_of_big_words+=1

                # if basic_word in middle_words or word in middle_words:
                #     continue
                # elif basic_word in high_words or word in high_words:
                #
                #     index = find_word_index(vocabulary['high'], word)
                #     if index == -1:
                #         high_index = high_index + 1
                #         vocabulary['high'].append((high_index, word, dicts[word]))
                #
                # else:
                #     print('超纲词:', word, basic_word)
                #     index = find_word_index(vocabulary['extra'], word)
                #     if index == -1:
                #         extra_index = extra_index + 1
                #         vocabulary['extra'].append((extra_index, word, dicts[word]))
    ##if method==1:（越大越难）
    difficulty1=number_of_words/number_of_sentence+number_of_big_words/number_of_words*100
    difficulty1=difficulty1*0.4
    ##elif method==2:(越大越容易)
    difficulty2 = number_of_words / number_of_sentence*1.015+ number_of_syllables/number_of_words*84.6
    difficulty2=206.835-difficulty2
    ##elif method==3:（越大越难）
    difficulty3=number_of_words / number_of_sentence*.0778+0.0455*number_of_syllables

    return (difficulty1, difficulty2, difficulty3)


####-------------------main-------------------
#
# middlePath = '初中词汇.txt'
# highPath = '高中词汇(含短语).txt'
# print('begin init dictionaries:')

# ----------------Readme说明-----------------
# 当前，初中词汇1400，高中2400，其中重复400，暂时以高中词汇为参考标准音标，
# 如果有不合适的过于简单的单词，从高中词汇表中移除即可
# for word in middle_wordset:
#     if word in high_wordset:
#         print(word)
#         high_wordset.remove(word)
# -------------------------------------------------
##计算单词的音节数
def syllable_count(word):
    word = word.lower()
    count = 0
    vowels = "aeiouy"
    if word[0] in vowels:
        count += 1
    for index in range(1, len(word)):
        if word[index] in vowels and word[index - 1] not in vowels:
            count += 1
    if word.endswith("e"):
        count -= 1
    if count == 0:
        count += 1
    return count


if __name__ == "__main__":
    with open('高考/初中词_基础词.json', 'r', encoding='utf-8') as f:
        middle_words = json.load(f)
    with open('高考/高中词汇2.json', 'r', encoding='utf-8') as f:
        high_words = json.load(f)
    with open('dicts_from_iciba.json', 'r', encoding='utf-8') as f:
        dicts = json.load(f)

    path_list = glob.glob('高考/提取的文章/*/*.docx')
    result = []

    for path in path_list:
        doc_src=docx.Document(path)
        doc, word_cout, vocabulary = get_new_article(doc_src, high_words, middle_words, dicts)

        difficulty = evaluate_difficulty(doc_src, vocabulary, 2)

        result.append((path, doc, word_cout, vocabulary, difficulty))



    # path = '高考/高考英语阅读1.docx'
    # outPath = '高考/abc.docx'
    # doc_src = docx.Document(path)
    # result=get_new_article(doc_src,  high_words, middle_words, dicts)

