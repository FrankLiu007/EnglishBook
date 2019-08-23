import docx
from docx.shared import Pt

# cambria 字体，Pt(16), 一行最多64个英文单词
maxNumber= 62
doc = docx.Document()
test =[
    ('city', ("['siti]", 'n. 城市, 市')),
    ('have', ('[hæv]', 'vt. 有, 怀有, 拿, 进行,aux. 已经')),
    ('good', ('[gud]', 'n. 善行, 好处, 利益,a. 好的, 优良的, 上等的, 愉快的, 有益的, 好心的, 慈善的, 虔诚的')),
    ('reason', ("['ri:zn]", 'n. 理由, 原因, 理智, 道理, 前提, 理性,vt. 说服, 推论, 辩论,vi. 推论,劝说, 思考')),
    ('for', ('[fɒ:]', 'prep. 为, 因为, 至于,conj. 因为')),
    ('being', ("['bi:iŋ]", 'n. 存在, 性质, 生命, 人, 生物, be的现在分词')),
    ('where', ('[hwєә]', 'ad. 在哪里,pron. 哪里,n. 地点')),
    ('they', ('[ðei]', 'pron. 他们, 它们')),
    ('are', ('[ɑ:]', 'be的现在时复数或第二人称单数')),
    ('like', ('[laik]', 'a. 相似的, 同样的,vt. 喜欢, 愿意, 想,vi. 喜欢, 希望,n. 爱好, 同样的人(或物),prep. 象, 如同,ad. 可能')),
    ('nearby', ("['niәbai]", 'a. 附近的, 近旁的,ad. 在附近, 近旁地,prep. 在...附近')),
    ('port', ('[pɒ:t]', 'n. 港口, 埠, 舱门, 避风港, 左舷, 炮眼, 姿势, 意义,vt. 左转舵, 持(枪),vi. 左转舵')),
    ('or', ('[ɒ:]', 'conj. 或, 或者')),
    ('river', ("['rivә]", 'n. 河, 江')),
    ('settle', ("['setl]", 'n. 有背长椅,vt. 决定, 整理, 安放, 使定居,使平静, 支付, 安排, 解决, 结算,vi. 停留, 下陷, 沉淀, 澄清, 安下心来, 结清, 定居, 安家')),
    ('in', ('[in]', 'prep. 在...期间, 在...之内, 处于...之中,从事于, 按照, 穿着,ad. 进入, 朝里, 在里面, 在屋里,a. 在里面的, 在朝的,n. 执政者, 交情')),
    ('these', ('[ði:z]', 'pron. 这些')),
    ('place', ('[pleis]','n. 地方, 地点, 位置, 住所, 座位, 地位, 处境, 特权, 空间, 余地, 职务, 位,vt. 放置, 寄予, 认出, 评定, 任命,vi. 名次列前')),
    ('because', ("[bi'kɒ:z]",'conj. 因为')),
    ('easy', ("['i:zi]", 'a. 容易的, 缓缓的, 舒适的, 从容的, 宽容的, 流畅的, 随便的, 自在的, 疲软的,ad. 容易地, 慢慢地')),
    ('to', ('[tu:]', 'prep. 到, 向, 趋于,ad. 向前')),
    ('get', ('[get]', 'vt. 得到, 获得, 变成, 使得, 收获, 接通, 抓住, 染上,vi. 到达, 成为, 变得,n. (网球等)救球, 生殖,幼兽')),
    ('and', ('[ænd]', 'conj. 和, 与')),
    ('natural', ("['nætʃәrәl]", 'n. 白痴,a. 自然的, 自然界的, 本能的, 天然的, 物质的, 正常的, 原始的, 自然数的')),
    ('suit', ('[sju:t, su:t]', 'n. 套装, 诉讼, 请求, 起诉, 套, 组,vt. 适合, 使适应,vi. 合适, 相称')),
    ('communicate', ("[kә'mju:nikeit]", 'vt.显露, 传达, 感染,vi. 通讯')),
    ('trade', ('[treid]', 'n. 贸易, 商业, 交易, 生意, 职业, 顾客, 信风,vi. 进行交易, 做买卖, 经商, 对换, 购物,vt. 用...进行交换')),
    ('new', ('[nju:]', 'a. 新的, 陌生的, 最近的, 不熟悉的')),
    ('example', ("[ig'zæmpl]", 'n. 例子, 样本, 实例')),
    ('near', ('[niә]', 'a. 近的,近亲的, 近似的,ad. 接近, 亲近,prep. 靠近, 近似于,vt. 接近, 走近,vi. 接近, 走近')),
    ('large', ('[lɑ:dʒ]', 'a. 大的, 大量的, 宽大的, 广博的,ad. 大大地,夸大地')),
    ('harbour', ('', 'n. 港, 避难所,vt. 庇护, 藏匿, (使)入港停泊,vi. 庇护, 藏匿, (使)入港停泊')),
    ('at', ('[æt]', 'prep. 在, 向, 对')),
    ('the', ('[ðә]', 'art. 那'))]

endList = [
    {'Cities': ('city', ("['siti]", 'n. 城市, 市'))},
    {'usually': None},
    {'have': ('have', ('[hæv]', 'vt. 有, 怀有, 拿, 进行,aux. 已经'))},
    {'a': None},
    {'good': ('good', ('[gud]', 'n. 善行, 好处, 利益,a. 好的, 优良的, 上等的, 愉快的, 有益的, 好心的, 慈善的, 虔诚的'))},
    {'reason': ('reason', ("['ri:zn]", 'n. 理由, 原因, 理智, 道理, 前提, 理性,vt. 说服, 推论, 辩论,vi. 推论, 劝说, 思考'))},
    {'for': ('for', ('[fɒ:]', 'prep. 为, 因为, 至于,conj. 因为'))},
    {'being': ('being', ("['bi:iŋ]", 'n. 存在, 性质, 生命, 人, 生物, be的现在分词'))},
    {'where': ('where', ('[hwєә]', 'ad. 在哪里,pron. 哪里,n. 地点'))},
    {'they': ('they', ('[ðei]', 'pron. 他们, 它们'))},
    {'are,': ('are', ('[ɑ:]', 'be的现在时复数或第二人称单数'))},
    {'like': ('like', ('[laik]', 'a. 相似的, 同样的,vt. 喜欢, 愿意, 想,vi. 喜欢, 希望,n. 爱好, 同样的人(或物),prep. 象, 如同,ad. 可能'))},
    {'a': None},
    {'nearby': ('nearby', ("['niәbai]", 'a. 附近的, 近旁的,ad. 在附近, 近旁地,prep. 在...附近'))},
    {'port': ('port', ('[pɒ:t]', 'n. 港口, 埠, 舱门, 避风港, 左舷, 炮眼, 姿势, 意义,vt.左转舵, 持(枪),vi. 左转舵'))},
    {'or': ('or', ('[ɒ:]', 'conj. 或, 或者'))},
    {'river': ('river', ("['rivә]", 'n. 河, 江'))},
    {'People': None},
    {'settle': ('settle', ("['setl]", 'n. 有背长椅 安下心来, 结清, 定居,安家'))},
    {'in': ('in', ('[in]', 'prep. 在...期间, 穿着,ad. 进入, 朝里, 在里面, 在屋里,a. 在里面的, 在朝的,n.执政者, 交情'))},
    {'these': ('these', ('[ði:z]', 'pron. 这些'))},
    {'places': ( 'place', ('[pleis]', 'n. 地方, 地点, 位置, 住所, 座位, 地位, 处境, 特权, 空间, 余地, 职务, 位,vt. 放置, 寄予, 认出, 评定, 任命,vi. 名次列前'))},
    {'because': ('because', ("[bi'kɒ:z]", 'conj. 因为'))},
    {'they': ('they', ('[ðei]', 'pron. 他们, 它们'))},
    {'are': ('are', ('[ɑ:]', 'be的现在时复数或第二人称单数'))},
    {'easy': ('easy', ("['i:zi]", 'a. 容易的, 缓缓的, 舒适的,从容的, 宽容的, 流畅的, 随便的, 自在的, 疲软的,ad. 容易地, 慢慢地'))},
    {'to': ('to', ('[tu:]', 'prep. 到, 向, 趋于,ad. 向前'))},
    {'get': ('get', ('[get]', 'vt. 得到, 获得, 变成, 使得, 收获, 接通, 抓住, 染上,vi. 到达, 成为, 变得,n. (网球等)救球, 生殖, 幼兽'))},
    {'to': ('to', ('[tu:]', 'prep. 到, 向, 趋于,ad. 向前'))},
    {'and': ('and', ('[ænd]', 'conj. 和, 与'))},
    {'naturally': ('natural', ("['nætʃәrәl]", 'n. 白痴,a. 自然的, 自然界的, 本能的, 天然的, 物质的, 正常的, 原始的, 自然数的'))},
    {'suited': ('suit', ('[sju:t, su:t]', 'n. 套装, 诉讼, 请求, 起诉, 套, 组,vt. 适合, 使适应,vi. 合适, 相称'))},
    {'to': ('to', ('[tu:]', 'prep. 到, 向, 趋于,ad. 向前'))},
    {'communications': ('communicate', ("[kә'mju:nikeit]", 'vt. 显露, 传达, 感染,vi. 通讯'))},
    {'and': ('and', ('[ænd]', 'conj. 和, 与'))},
    {'trade': ('trade', ('[treid]', 'n. 贸易, 商业, 交易, 生意, 职业, 顾客, 信风,vi. 进行交易, 做买卖,经商, 对换, 购物,vt. 用...进行交换'))},
    {'New': ('new', ('[nju:]', 'a. 新的, 陌生的, 最近的, 不熟悉的'))},
    {'York': None},
    {'City,': ('city', ("['siti]", 'n. 城市, 市'))},
    {'for': ('for', ('[fɒ:]', 'prep. 为, 因为, 至于,conj. 因为'))},
    {'example,': ('example', ("[ig'zæmpl]", 'n. 例子, 样本, 实例'))},
    {'is': None},
    {'near': ('near', ('[niә]', 'a. 近的, 近亲的, 近似的,ad. 接近, 亲近,prep. 靠近, 近似于,vt. 接近, 走近,vi. 接近, 走近'))},
    {'a': None},
    {'large': ('large', ('[lɑ:dʒ]', 'a. 大的, 大量的, 宽大的, 广博的,ad. 大大地, 夸大地'))},
    {'harbour': ('harbour', ('', 'n. 港, 避难所,vt. 庇护, 藏匿, (使)入港停泊,vi. 庇护, 藏匿, (使)入港停泊'))},
    {'at': ('at', ('[æt]', 'prep. 在, 向, 对'))}, {'the': ('the', ('[ðә]', 'art. 那'))},
    {'mouth': ('mouth', ('[mauθ]', 'n. 嘴, 口, 口腔, 口状物,vi. 装腔作势说话, 做鬼脸,vt. 说出, 做作地说'))},
    {'of': ('of', ('[ɒv]', 'prep. 的, 属于'))},
    {'the': ('the', ('[ðә]', 'art. 那'))},
    {'Hudson': None},
    {'River': ('river', ("['rivә]", 'n. 河, 江'))},
    {'Over': ('over', ("['әuvә]", 'ad. 结束, 越过, 从头到尾,prep. 在...之上, 遍于...之上, 越过,a. 上面的,vt. 越过'))},
    {'300': None},
    {'years': ('year', ('[jiә]', 'n. 年, 年度, 年龄'))},
    {'its': ('its', ('[its]', 'pron. 它的'))},
    {'population': ('population', ("[,pɒpju'leiʃәn]", 'n. 人口, 人口数'))},
    {'grew': ('grow', ('[grәu]', 'vt. 种植, 使长满,vi. 生长,变成, 发展'))},
    {'gradually': ('grade', ('[greid]', 'n. 等级, 年级, 阶段, 成绩, 程度, 坡度, 斜坡,vt. 分等, 分级, 评分,vi. 属于某等级, 逐渐变化'))},
    {'from': ('from', ('[frɒm]', 'prep. 从, 来自, 根据'))},
    {'800': None}, {'people': None},
    {'to': ('to', ('[tu:]', 'prep. 到, 向, 趋于,ad. 向前'))},
    {'8': None},
    {'million': ('million', ("['miljәn]", 'n. 百万, 无数,num. 百万'))},
    {'But': ('but', ('[bʌt]', 'prep. 除了,conj. 但是,ad. 仅仅'))},
    {'not': ('not', ('[nɒt]', 'ad. 不, 非, 未'))},
    {'all': ('all', ('[ɒ:l]', 'a. 所有的, 全部的, 一切的,ad. 全部, 全然,pron. 全部,n. 全部'))},
    {'cities': ('city', ("['siti]", 'n. 城市, 市'))},
    {'develop': ('develop', ("[di'velәp]", 'vt. 发展, 使发达, 进步, 洗印, 显影,vi. 发展, 生长'))},
    {'slowly': None},
    {'over': ('over', ("['әuvә]", 'ad. 结束, 越过, 从头到尾,prep. 在...之上, 遍于...之上, 越过,a. 上面的,vt. 越过'))},
    {'a': None},
    {'long': ('long', ('[lɒŋ]', 'a. 长的, 长久的, 冗长的, 做多头的,vi. 渴望, 热望, 极想,ad. 长久, 始终,n. 长时间, 长信号, 长整型'))},
    {'period': ('period', ("['piәriәd]", 'n.时期, 节段, 节, 句点, 学时, 周期,a. 当时特有的, 过去某段时期的,interj. 就是这话, 就是这么回事'))},
    {'of': ('of', ('[ɒv]', 'prep. 的, 属于'))},
    {'time': ('time', ( '[taim]', 'n. 时间, 时侯, 时机, 时期, 期限, 次数, 节拍, 暂停, 规定时间,vt. 测定...的时间, 记录...的时间, 计时, 定时,a. 时间的, 记时的,定时的, 定期的, 分期的'))}
]


def complete_blank(word , say):
    pL = len(word)
    sL = len(say)
    lL = ((pL+1)*3-sL)/2
    lL = round(lL)

    if lL <0:
        return say

    i = 0
    blankWord = ''
    while i < lL:
        blankWord = blankWord+'.'
        i+=1

    say = blankWord + say + blankWord
    return say

def get_a_line(wordList):
    showWord = ''
    for word in wordList:
        strlen = len(showWord)
        showWord = showWord + word +' '
        if(len(showWord)>maxNumber):
            showWord = showWord[:strlen]
            set_other_line(showWord)
            showWord = word+' '

def set_other_line(showWord):
    showWordList = showWord.split(' ')
    sayWordList = []
    for word in showWordList:
        for item in endList:
            if word in item:
                wordMes = item[word]
                if wordMes is not None:
                    say = wordMes[1][0]
                else:
                    say = ''
                sayWord = complete_blank(word, say)
                sayWordList.append(sayWord)
                break
    write_a_line(showWordList , sayWordList)

def write_a_line(showArr , sayArr):
    showTxt = ''
    sayTxt = ''
    for show in showArr:
        showTxt = showTxt+show+' '
    for say in sayArr:
        sayTxt = sayTxt+say+' '
    write_to_word(sayTxt , showTxt)

def write_to_word(txt1 , txt2):
    paragraph1 = doc.add_paragraph()
    run1 = paragraph1.add_run(txt1)
    run1.font.size = Pt(8)
    paragraph2 = doc.add_paragraph()
    run2 = paragraph2.add_run(txt2)
    run2.font.size = Pt(16)

path = '高考英语阅读.txt'
with open(path, 'r', encoding='UTF-8') as f:
    for ll in f.readlines():
        ll = ll.split(' ')
        get_a_line(ll)
    doc.save('test1.docx')
