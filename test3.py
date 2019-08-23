from docx import Document
from docx.shared import Pt

document = Document()
p = document.add_paragraph()

run1 = p.add_run('Cities usually have a good reason for being where they are, like')
run1.font.size = Pt(16)
run2 = p.add_run('Cities usually have a good reason for being where they are, like')
run2.font.size = Pt(8)

document.save('test3.docx')